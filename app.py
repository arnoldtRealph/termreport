import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Streamlit Config
st.set_page_config(page_title="📊 Learner Performance Dashboard", layout="wide")

# Custom Styling with Explicit Colors
st.markdown("""
    <style>
        .stApp {
            background-color: #F4F6F9;
            color: #333333;
            font-family: Arial, sans-serif;
        }
        .custom-title {
            font-size: 30px;
            color: #003366;
        }
        .welcome-text {
            font-size: 20px;
            color: #4CAF50;
            text-align: center;
            margin-bottom: 20px;
        }
        .footer {
            text-align: center;
            font-size: 12px;
            color: #003366;
            margin-top: 20px;
        }
        .sidebar .sidebar-content {
            background-color: #E8ECEF;
            padding: 10px;
        }
        .sidebar h3 {
            font-size: 18px;
            color: #003366;
        }
        h2, h3 {
            color: #003366;
        }
        .stButton>button {
            background-color: #003366;
            color: white;
        }
        .larger-instructions .stAlert {
            font-size: 20px;  /* Larger font size for instructions */
        }
    </style>
""", unsafe_allow_html=True)

# Sidebar Configuration
st.sidebar.title("📊 Dashboard Options")
st.sidebar.markdown("### File Upload")
uploaded_file = st.sidebar.file_uploader("Upload your Excel file", type=["xlsx"])

st.sidebar.markdown("### Navigation")
sections = {
    "Learner Overview": "overview",
    "Question Analysis": "charts",
    "Insights and Recommendations": "recommendations",
    "Download Full Report": "download"
}
for section_name, section_id in sections.items():
    st.sidebar.markdown(f'<a href="#{section_id}" style="text-decoration: none; color: #003366;">{section_name}</a>', unsafe_allow_html=True)

st.sidebar.markdown("### Chart Settings")
chart_options = ["Vertical Bar", "Scatter Plot", "Box Plot"]
selected_chart = st.sidebar.selectbox("Select chart type for Total Marks per Learner", chart_options, index=0)

# Main content
st.markdown("<h1 class='custom-title' style='text-align: center;'>Saul Damon High School</h1>", unsafe_allow_html=True)
st.markdown("<div class='welcome-text'>**Welkom julle.** Gebruik hierdie tool wat ek ontwerp het om handige insig te kry vanuit leerders se toetse en take.</div>", unsafe_allow_html=True)

# Improved Instructions with Larger Text
st.markdown('<div class="larger-instructions">', unsafe_allow_html=True)
st.info("""
### Instructions
1. Gebruik die excel templaat wat ek aan julle voorsien het om die punte per vraag analise te doen.  
2. Upload dan die excel file in die sidebar hier langsaan.  
3. Download dan die Word dokument en stoor in die masterfile.
""", icon="ℹ️")
st.markdown('</div>', unsafe_allow_html=True)

if uploaded_file:
    # Read the Excel file
    df_raw = pd.read_excel(uploaded_file, header=None)
    header_row = next((i for i, row in df_raw.iterrows() if any("name of learner" in str(cell).lower() for cell in row)), None)
    if header_row is not None:
        df = pd.read_excel(uploaded_file, skiprows=header_row)
        df.columns = df.columns.str.strip()
        name_col = next((col for col in df.columns if "name of learner" in col.lower()), None)
        if name_col:
            name_col_idx = df.columns.get_loc(name_col)
            question_cols = [col for col in df.columns[name_col_idx + 1:] if df[col].dtype in ['int64', 'float64']]
            if question_cols:
                df[question_cols] = df[question_cols].fillna(0)
                df['Total'] = df[question_cols].sum(axis=1)
                max_total = df['Total'].max()
                df['Percentage'] = (df['Total'] / max_total * 100) if max_total > 0 else 0
            else:
                st.error("No numeric question columns found after 'NAME OF LEARNER'.")
                st.stop()
        else:
            st.error("Could not find 'NAME OF LEARNER' column.")
            st.stop()
    else:
        st.error("Could not locate 'NAME OF LEARNER' row.")
        st.stop()

    # Define function to save chart to buffer
    def save_chart_to_buffer(fig):
        buf = BytesIO()
        fig.savefig(buf, format="png")
        buf.seek(0)
        return buf

    # Learner Overview
    st.markdown(f'<a name="overview"></a>', unsafe_allow_html=True)
    st.subheader("Results DASHBOARD")
    st.write("Upload your class mark sheet to generate insightful analysis.")
    
    st.subheader("📊 Average Percentage per Learner")
    avg_fig, ax1 = plt.subplots(figsize=(8, max(5, len(df) * 0.3)))
    bars = ax1.barh(df[name_col], df['Percentage'], color=sns.color_palette("Blues_d", len(df)))
    for bar in bars:
        bar.set_edgecolor('#003366')
        bar.set_linewidth(1)
    ax1.set_title("Percentage per Learner", color='#003366', pad=20)
    ax1.set_xlabel("Percentage (%)")
    ax1.set_ylabel("Learner")
    ax1.invert_yaxis()
    ax1.grid(True, axis='x', linestyle='--', alpha=0.7)
    plt.tight_layout()
    st.pyplot(avg_fig)

    st.subheader("📊 Total Marks per Learner")
    if selected_chart == "Vertical Bar":
        total_fig, ax1 = plt.subplots(figsize=(10, 6))
        sns.barplot(x=df[name_col], y=df['Total'], ax=ax1, palette="Blues_d")
        ax1.set_title("Total Marks per Learner", color='#003366', pad=20)
        ax1.set_xlabel("Learner")
        ax1.set_ylabel("Total Marks")
        ax1.set_xticklabels(ax1.get_xticklabels(), rotation=90, ha='center')
        ax1.grid(True, axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        st.pyplot(total_fig)
    elif selected_chart == "Scatter Plot":
        total_fig, ax1 = plt.subplots(figsize=(10, 6))
        ax1.scatter(df[name_col], df['Total'], color='#003366', edgecolor='black', s=100)
        ax1.set_title("Total Marks per Learner", color='#003366', pad=20)
        ax1.set_xlabel("Learner")
        ax1.set_ylabel("Total Marks")
        ax1.set_xticklabels(df[name_col], rotation=60, ha='right')
        ax1.grid(True, linestyle='--', alpha=0.7)
        plt.tight_layout()
        st.pyplot(total_fig)
    elif selected_chart == "Box Plot":
        total_fig, ax1 = plt.subplots(figsize=(10, 6))
        sns.boxplot(y=df['Total'], ax=ax1, color='#003366', boxprops=dict(edgecolor='black'))
        ax1.set_title("Distribution of Total Marks", color='#003366', pad=20)
        ax1.set_ylabel("Total Marks")
        ax1.grid(True, axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        st.pyplot(total_fig)

    # Stacked Bar Chart with Adjusted Names
    st.subheader("📊 Marks Breakdown by Learner and Question")
    stacked_fig, ax = plt.subplots(figsize=(10, 6))
    df.set_index(name_col)[question_cols].plot(kind='bar', stacked=True, ax=ax, colormap="Paired")
    ax.set_title("Marks Breakdown by Learner and Question", color='#003366', pad=20)
    ax.set_xlabel("Learner")
    ax.set_ylabel("Marks")
    ax.set_xticklabels(ax.get_xticklabels(), rotation=90, ha='center')  # Match Total Marks chart
    ax.legend(title="Questions", bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.tight_layout()
    st.pyplot(stacked_fig)

    # Question Analysis
    st.markdown(f'<a name="charts"></a>', unsafe_allow_html=True)
    st.subheader("📊 Question Analysis")
    active_questions = [q for q in question_cols if df[q].sum() > 0]
    
    if active_questions:
        st.subheader("Average Marks per Question")
        q_avg_fig, ax_q = plt.subplots(figsize=(8, 5))
        active_means = df[active_questions].mean()
        sns.barplot(x=active_means.index, y=active_means.values, ax=ax_q, palette="Greens_d")
        ax_q.set_title("Average Marks per Question", color='#003366')
        ax_q.set_xlabel("Question")
        ax_q.set_ylabel("Average Mark")
        ax_q.set_xticklabels(ax_q.get_xticklabels(), rotation=45, ha='right')
        st.pyplot(q_avg_fig)

        st.subheader("Distribution per Question")
        for question in active_questions:
            pie_fig, ax2 = plt.subplots(figsize=(6, 6))
            question_data = df[question].value_counts()
            total = question_data.sum()
            percentages = [(count / total * 100) for count in question_data]
            wedges, texts, autotexts = ax2.pie(
                question_data,
                startangle=90,
                colors=sns.color_palette("Set2", len(question_data)),
                autopct='%1.1f%%',
                pctdistance=0.85,
                wedgeprops={'edgecolor': 'white', 'linewidth': 1}
            )
            ax2.axis('equal')
            ax2.set_title(f"Distribution of Marks for {question}", color='#003366', pad=20)
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontsize(10)
                autotext.set_weight('bold')
            ax2.legend(
                labels=[f"{mark}: {pct:.1f}%" for mark, pct in zip(question_data.index, percentages)],
                title="Marks (Percentage)",
                loc="center left",
                bbox_to_anchor=(1, 0, 0.5, 1)
            )
            plt.tight_layout()
            st.pyplot(pie_fig)
    else:
        st.write("No questions have marks entered yet.")

    # Improved Insights and Recommendations with Table
    st.markdown(f'<a name="recommendations"></a>', unsafe_allow_html=True)
    st.subheader("📝 Insights and Recommendations")
    
    insights = []
    recommendations = []
    learner_performance = {"Needs Improvement": [], "Average": [], "Excelling": []}

    avg_percentage = df['Percentage'].mean()
    insights.append(f"The class average percentage is **{avg_percentage:.2f}%**.")

    if active_questions:
        active_means = df[active_questions].mean()
        avg_question_mean = active_means.mean()
        weak_questions = active_means[active_means < avg_question_mean * 0.7].index.tolist()
        if weak_questions:
            insights.append(f"Weak questions (below 70% of mean {avg_question_mean:.1f}): **{', '.join(weak_questions)}**.")

        # Per-Learner Analysis with Tiered Summary
        for learner in df[name_col]:
            learner_data = df[df[name_col] == learner]
            percentage = learner_data['Percentage'].iloc[0]
            good_questions = [q for q in active_questions if learner_data[q].iloc[0] > active_means[q]]
            bad_questions = [q for q in active_questions if learner_data[q].iloc[0] < active_means[q]]
            
            if percentage < 40:
                tier = "Needs Improvement"
            elif percentage < 70:
                tier = "Average"
            else:
                tier = "Excelling"
            
            learner_performance[tier].append({
                "Learner": learner,
                "Percentage": f"{percentage:.1f}%",
                "Strong Questions": ', '.join(good_questions) or "None",
                "Weak Questions": ', '.join(bad_questions) or "None"
            })

    # Insights as Tables
    st.markdown("### Insights")
    st.markdown(f"- The class average percentage is {avg_percentage:.2f}%.")
    if active_questions and weak_questions:
        st.markdown(f"- Weak questions (below 70% of mean {avg_question_mean:.1f}): {', '.join(weak_questions)}.")

    for tier, learners in learner_performance.items():
        if learners:
            st.markdown(f"#### {tier} ({len(learners)} learners)")
            st.table(pd.DataFrame(learners))

    # Recommendations
    recommendations.append("### Recommendations")
    if learner_performance["Needs Improvement"]:
        recommendations.append("- Organize remedial sessions for 'Needs Improvement' learners.")
    if learner_performance["Excelling"]:
        recommendations.append("- Provide advanced challenges for 'Excelling' learners.")
    if active_questions:
        weakest_question = active_means.idxmin()
        strongest_question = active_means.idxmax()
        recommendations.append(f"- Focus class revision on **{weakest_question}** (lowest average: {active_means[weakest_question]:.1f}).")
        recommendations.append(f"- Note strong performance in **{strongest_question}** (highest average: {active_means[strongest_question]:.1f}).")

    for recommendation in recommendations:
        st.markdown(recommendation)

    # Download Full Report
    st.markdown(f'<a name="download"></a>', unsafe_allow_html=True)
    st.subheader("📥 Download Full Report (Word Document)")
    if st.button("Download Full Report"):
        doc = Document()

        # Title Page
        doc.add_heading('Learner Performance Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Saul Damon High School', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Generated on April 05, 2025', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Performance Overview
        doc.add_heading('Performance Overview', level=1)
        doc.add_picture(save_chart_to_buffer(avg_fig), width=Inches(5.5))
        doc.add_paragraph("Figure 1: Average Percentage per Learner", style='Caption')
        doc.add_picture(save_chart_to_buffer(total_fig), width=Inches(5.5))
        doc.add_paragraph(f"Figure 2: Total Marks per Learner ({selected_chart})", style='Caption')
        doc.add_picture(save_chart_to_buffer(stacked_fig), width=Inches(5.5))
        doc.add_paragraph("Figure 3: Marks Breakdown by Learner and Question", style='Caption')

        # Question Analysis
        doc.add_heading('Question Analysis', level=1)
        if active_questions:
            doc.add_picture(save_chart_to_buffer(q_avg_fig), width=Inches(5.5))
            doc.add_paragraph("Figure 4: Average Marks per Question", style='Caption')
            for i, question in enumerate(active_questions, 5):
                pie_fig, _ = plt.subplots(figsize=(6, 6))
                question_data = df[question].value_counts()
                total = question_data.sum()
                percentages = [(count / total * 100) for count in question_data]
                plt.pie(
                    question_data,
                    startangle=90,
                    colors=sns.color_palette("Set2", len(question_data)),
                    autopct='%1.1f%%',
                    pctdistance=0.85,
                    wedgeprops={'edgecolor': 'white', 'linewidth': 1}
                )
                plt.title(f"Distribution of Marks for {question}", color='#003366', pad=20)
                plt.legend(
                    labels=[f"{mark}: {pct:.1f}%" for mark, pct in zip(question_data.index, percentages)],
                    title="Marks (Percentage)",
                    loc="center left",
                    bbox_to_anchor=(1, 0, 0.5, 1)
                )
                plt.axis('equal')
                plt.tight_layout()
                doc.add_picture(save_chart_to_buffer(pie_fig), width=Inches(5.5))
                doc.add_paragraph(f"Figure {i}: Distribution of Marks for {question}", style='Caption')

        # Insights and Recommendations
        doc.add_heading('Insights and Recommendations', level=1)
        doc.add_paragraph("Insights", style='Heading 2')
        doc.add_paragraph(f"The class average percentage is {avg_percentage:.2f}%.", style='List Bullet')
        if active_questions and weak_questions:
            doc.add_paragraph(f"Weak questions (below 70% of mean {avg_question_mean:.1f}): {', '.join(weak_questions)}.", style='List Bullet')

        for tier, learners in learner_performance.items():
            if learners:
                doc.add_paragraph(f"{tier} ({len(learners)} learners)", style='Heading 3')
                table = doc.add_table(rows=len(learners) + 1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Learner'
                hdr_cells[1].text = 'Percentage'
                hdr_cells[2].text = 'Strong Questions'
                hdr_cells[3].text = 'Weak Questions'
                for i, learner in enumerate(learners, 1):
                    row_cells = table.rows[i].cells
                    row_cells[0].text = learner['Learner']
                    row_cells[1].text = learner['Percentage']
                    row_cells[2].text = learner['Strong Questions']
                    row_cells[3].text = learner['Weak Questions']
                for row in table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph("Recommendations", style='Heading 2')
        for recommendation in recommendations:
            if recommendation.startswith("###"):
                doc.add_paragraph(recommendation.replace("### ", ""), style='Heading 3')
            else:
                doc.add_paragraph(recommendation.replace('**', ''), style='List Bullet')

        # Save and Download
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        st.download_button(
            label="Download Full Report (Word Document)",
            data=doc_stream,
            file_name="learner_performance_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Footer
st.markdown(f'<div class="footer">Designed by Mr AR Visagie</div>', unsafe_allow_html=True)